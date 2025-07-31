from fastapi import APIRouter, Query, UploadFile, File, Form, Depends, HTTPException, status
from pydantic import BaseModel
from dotenv import load_dotenv
import os
from google import genai
import markdown2
from fastapi.responses import JSONResponse
from typing import List
import shutil
import tempfile
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain.vectorstores import Chroma
from langchain.chains import RetrievalQA
import docx
import json # Import for JSON processing
from app.api.userapi import get_current_user
from langchain.docstore.document import Document # Import for creating LangChain Document objects

# Load environment variables
load_dotenv()

# Setup Gemini Clients with two API keys
api_key1 = os.getenv("GEMINI_API_KEY1")
api_key2 = os.getenv("GEMINI_API_KEY2")

# Initialize Gemini clients
# It's good practice to handle cases where API keys might be missing
if api_key1:
    client1 = genai.Client(api_key=api_key1)
else:
    client1 = None
    print("Warning: GEMINI_API_KEY1 not found in environment variables.")

if api_key2:
    client2 = genai.Client(api_key=api_key2)
else:
    client2 = None
    print("Warning: GEMINI_API_KEY2 not found in environment variables.")


router = APIRouter()

class ChatRequest(BaseModel):
    prompt: str

class RAGRequest(BaseModel):
    query: str

class GoalItem(BaseModel):
    goal: str
    checked: bool

goals_data = []

# Store vectorstore and retriever in memory for demo (per session)
# In a production environment, you might want to persist this or manage it per user.
vectorstore = None
vector_index = None
rag_model = None

# Helper to extract text from DOCX
def extract_docx_text(docx_path):
    """
    Extracts plain text from a DOCX file.
    """
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

# Helper to extract text from JSON
def extract_json_text(json_path):
    """
    Extracts and formats JSON content into a readable string.
    This helps the LLM understand the structure.
    """
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    # Use json.dumps with indent for pretty printing, making it more readable for the LLM
    return json.dumps(data, indent=2)

# Helper to extract text from Python files (or any plain text file)
def extract_plain_text(file_path):
    """
    Extracts plain text content from a file.
    Suitable for .py, .txt, .md files.
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    return content

@router.post("/api/upload")
async def upload_file(file: UploadFile = File(...), current_user: dict = Depends(get_current_user)):
    """
    Uploads a file (PDF, DOCX, PY, JSON), processes its content,
    and sets up the RAG model for querying.
    """
    global vectorstore, vector_index, rag_model
    
    # Validate file type
    allowed_extensions = ('.pdf', '.docx', '.py', '.json')
    if not file.filename.lower().endswith(allowed_extensions):
        raise HTTPException(status_code=400, detail=f"Unsupported file type. Only {', '.join(allowed_extensions)} allowed.")
    
    # Validate file size (max 10MB)
    max_size = 10 * 1024 * 1024  # 10MB
    if file.size and file.size > max_size:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 10MB.")
    
    tmp_path = None
    try:
        # Save uploaded file to temp
        # Using file.filename.split('.')[-1] to get the actual extension for suffix
        file_extension = "." + file.filename.split('.')[-1]
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp:
            shutil.copyfileobj(file.file, tmp)
            tmp_path = tmp.name
        
        pages = []
        if file.filename.lower().endswith(".pdf"):
            try:
                loader = PyPDFLoader(tmp_path)
                pages = loader.load()
            except ImportError as e:
                if "pypdf" in str(e):
                    raise HTTPException(
                        status_code=500, 
                        detail="PDF processing library (pypdf) not installed. Please contact administrator."
                    )
                else:
                    raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading PDF file: {str(e)}")
        elif file.filename.lower().endswith(".docx"):
            try:
                text = extract_docx_text(tmp_path)
                # Wrap as LangChain Document
                pages = [Document(page_content=text, metadata={"source": file.filename})]
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading DOCX file: {str(e)}")
        elif file.filename.lower().endswith(".py"):
            try:
                text = extract_plain_text(tmp_path)
                # Wrap as LangChain Document, adding metadata for source
                pages = [Document(page_content=text, metadata={"source": file.filename, "file_type": "python_code"})]
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading Python file: {str(e)}")
        elif file.filename.lower().endswith(".json"):
            try:
                text = extract_json_text(tmp_path)
                # Wrap as LangChain Document, adding metadata for source
                pages = [Document(page_content=text, metadata={"source": file.filename, "file_type": "json_data"})]
            except Exception as e:
                raise HTTPException(status_code=400, detail=f"Error reading JSON file: {str(e)}")
        
        if not pages:
            raise HTTPException(status_code=400, detail="No readable content found in the file.")
        
        # Split into chunks
        # For code and JSON, a smaller chunk_size and larger overlap might be beneficial
        # to keep logical units together or provide more context.
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200, separators=['\n\n', '\n', ' ', ''])
        chunks = splitter.split_documents(pages)
        
        if not chunks:
            raise HTTPException(status_code=400, detail="No processable chunks found in the file.")
        
        # Embeddings and vectorstore
        api_key = api_key1 or api_key2
        if not api_key:
            raise HTTPException(status_code=500, detail="API key not configured.")
        
        # Initialize embeddings with the selected API key
        embeddings = GoogleGenerativeAIEmbeddings(model='models/embedding-001', google_api_key=api_key)
        
        # Create or update the vectorstore
        # For simplicity, this overwrites the previous vectorstore.
        # For multi-file support where you want to query across all uploaded files,
        # you'd append to an existing vectorstore or manage multiple ones.
        vectorstore = Chroma.from_documents(documents=chunks, embedding=embeddings)
        vector_index = vectorstore.as_retriever(search_kwargs={'k': 4}) # Reduced k for potentially more focused retrieval
        
        # RAG model
        # Using gemini-2.0-flash for faster responses
        llm_model = ChatGoogleGenerativeAI(model="gemini-2.0-flash", google_api_key=api_key, temperature=0.2)
        
        # Define a custom prompt for the RAG model to guide it on "approach" and "technology"
        # This is implicitly handled by the LLM's capabilities and the quality of retrieved chunks.
        # LangChain's RetrievalQA chain is designed to integrate context.
        # The LLM itself will be responsible for interpreting "approach" and "technology"
        # based on the provided code/JSON snippets.
        
        rag_model = RetrievalQA.from_chain_type(
            llm=llm_model,
            chain_type="stuff", # "stuff" puts all retrieved documents into one prompt
            retriever=vector_index,
            return_source_documents=True
        )
        
        return {"message": "File uploaded and processed successfully.", "chunks_count": len(chunks)}
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Internal server error during file processing: {str(e)}")
    finally:
        # Clean up temporary file
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.unlink(tmp_path)
            except Exception as e:
                print(f"Error deleting temporary file {tmp_path}: {e}") # Log error but don't block response

@router.post("/api/rag_chat")
async def rag_chat(req: RAGRequest, current_user: dict = Depends(get_current_user)):
    """
    Answers questions based on the content of the uploaded document using RAG.
    """
    global rag_model
    if rag_model is None:
        raise HTTPException(status_code=400, detail="No document uploaded yet. Please upload a file first.")
    
    if not req.query or not req.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty.")
    
    try:
        # The RetrievalQA chain handles fetching relevant documents and passing them to the LLM.
        # The LLM's ability to answer "approach" and "technology" questions will depend on
        # how well the retrieved chunks represent those concepts in the code/JSON.
        response = rag_model({"query": req.query})
        answer = response["result"]
        
        # Optionally, include sources
        sources = []
        for doc in response.get("source_documents", []):
            source_info = doc.metadata.get("source", "Unknown Source")
            # If it's a code or JSON file, you might want to add more specific info
            if doc.metadata.get("file_type") == "python_code":
                sources.append(f"Code file: {source_info}")
            elif doc.metadata.get("file_type") == "json_data":
                sources.append(f"JSON data: {source_info}")
            else:
                sources.append(source_info)

        return {"answer": answer, "sources": sources}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing RAG query: {str(e)}")

async def generate_reply(prompt: str, client):
    """
    Generates a reply using the specified Gemini client.
    Includes exponential backoff for API calls.
    """
    if not client:
        raise ValueError("Gemini client not initialized. Check API keys.")

    retries = 0
    max_retries = 5
    while retries < max_retries:
        try:
            response = client.models.generate_content(
                model="gemini-2.0-flash",
                contents=prompt
            )
            return response.text
        except Exception as e:
            # Check for rate limit or transient errors (e.g., 503)
            if "503" in str(e) or "Service temporarily unavailable" in str(e) or "rate limit" in str(e).lower():
                retries += 1
                wait_time = 2 ** retries # Exponential backoff
                print(f"API call failed (retry {retries}/{max_retries}): {e}. Retrying in {wait_time} seconds...")
                await asyncio.sleep(wait_time) # Use asyncio.sleep in async function
            else:
                raise # Re-raise other unexpected errors
    raise Exception(f"Failed to generate reply after {max_retries} retries.")


@router.post("/api/chat")
async def chat(req: ChatRequest, as_markdown: bool = Query(False)):
    """
    Handles general chat requests using Gemini API with fallback.
    """
    # Ensure asyncio is imported for sleep
    import asyncio 

    if not client1 and not client2:
        raise HTTPException(status_code=500, detail="No Gemini API keys configured. Cannot perform chat.")

    try:
        reply_text = ""
        # Try with client1 first
        if client1:
            try:
                reply_text = await generate_reply(req.prompt, client1)
            except Exception as e1:
                print(f"Client1 failed: {e1}. Trying client2...")
                # If client1 fails, try client2
                if client2:
                    try:
                        reply_text = await generate_reply(req.prompt, client2)
                    except Exception as e2:
                        print(f"Client2 also failed: {e2}.")
                        return {"reply": "Service is temporarily busy. Please try again later."}
                else:
                    return {"reply": "Service is temporarily busy. Please try again later (only one API key available)."}
        elif client2: # If client1 was not configured, directly try client2
            try:
                reply_text = await generate_reply(req.prompt, client2)
            except Exception as e2:
                print(f"Client2 failed: {e2}.")
                return {"reply": "Service is temporarily busy. Please try again later."}
        else:
            return {"reply": "No active Gemini clients available."}

        if as_markdown:
            reply_text = markdown2.markdown(reply_text)
        return {"reply": reply_text}
    except Exception as e:
        return {"reply": f"An unexpected error occurred during chat: {str(e)}"}

@router.post("/api/goals")
async def save_goals(goals: list[GoalItem]):
    """
    Saves a list of goals. (Unchanged from original)
    """
    global goals_data
    goals_data = goals
    return {"message": "Goals saved successfully"}

@router.get("/api/goals")
async def get_goals():
    """
    Retrieves the saved list of goals. (Unchanged from original)
    """
    return goals_data

@router.get("/api/performance")
async def get_performance():
    """
    Calculates and returns goal completion performance. (Unchanged from original)
    """
    total = len(goals_data)
    completed = sum(1 for g in goals_data if g.checked)
    percent = (completed / total) * 100 if total > 0 else 0
    return {
        "total": total,
        "completed": completed,
        "percent": percent
    }

@router.get("/api/upload/health")
async def upload_health_check():
    """
    Health check for upload functionality, now including PY and JSON.
    """
    try:
        # Check if required dependencies are available
        import langchain
        import docx
        import tempfile
        import os
        import json # Check for json module
        
        # Check PDF dependencies specifically
        try:
            import pypdf
            pdf_status = "available"
        except ImportError:
            pdf_status = "missing pypdf"
        
        # PyPDF2 is often used with older PyPDFLoader versions or as an alternative.
        # If pypdf is preferred, PyPDF2 might not be strictly necessary.
        try:
            import PyPDF2
            pypdf2_status = "available"
        except ImportError:
            pypdf2_status = "missing PyPDF2"
        
        # Check if API keys are configured
        api_key = api_key1 or api_key2
        if not api_key:
            return {"status": "error", "message": "API keys not configured"}
        
        return {
            "status": "healthy" if pdf_status == "available" else "warning",
            "message": "Upload service is ready" if pdf_status == "available" else "PDF support may be limited",
            "supported_formats": ["PDF", "DOCX", "PY", "JSON"], # Updated supported formats
            "max_file_size": "10MB",
            "dependencies": {
                "pypdf": pdf_status,
                "PyPDF2": pypdf2_status, # Still useful to show
                "python-docx": "available",
                "json_module": "available" # Indicate json module is available
            }
        }
    except ImportError as e:
        return {"status": "error", "message": f"Missing core dependency: {str(e)}. Please ensure all required libraries are installed."}
    except Exception as e:
        return {"status": "error", "message": f"Service error during health check: {str(e)}"}

